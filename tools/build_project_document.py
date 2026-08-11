from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / 'docs' / 'AgriLink_Project_Documentation.docx'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GREY = 'F2F4F7'
INK = '17324D'
MUTED = '5B6573'
GREEN = '2F6B45'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    indent = tbl_pr.first_child_found_in('w:tblInd')
    if indent is None:
        indent = OxmlElement('w:tblInd')
        tbl_pr.append(indent)
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        if index < len(grid.gridCol_lst):
            grid.gridCol_lst[index].set(qn('w:w'), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in('w:tcMar')
            if margins is None:
                margins = OxmlElement('w:tcMar')
                tc_pr.append(margins)
            for side in ('top', 'start', 'bottom', 'end'):
                node = margins.find(qn(f'w:{side}'))
                if node is None:
                    node = OxmlElement(f'w:{side}')
                    margins.append(node)
                node.set(qn('w:w'), '80' if side in ('top', 'bottom') else '120')
                node.set(qn('w:type'), 'dxa')


def set_font(run, name='Calibri', size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, **kwargs):
    return set_font(paragraph.add_run(text), **kwargs)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        add_text(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        add_text(p, item)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=INK)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_text(p, header, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            add_text(p, str(value))
            if len(row) % 2 == 0 and len(rows) > 4 and row == rows[-1]:
                pass
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F4F6F9')
    p = cell.paragraphs[0]
    add_text(p, title + ' ', bold=True, color=DARK_BLUE)
    add_text(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F4F6F9')
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run('\n')
        add_text(p, line, name='Consolas', size=9, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ('Title', 26, INK, 0, 10),
        ('Subtitle', 12, MUTED, 0, 16),
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 14, 7),
        ('Heading 3', 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name not in ('Subtitle',)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for list_name in ('List Bullet', 'List Number'):
        style = styles[list_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, 'AGRILINK | Technical Project Documentation', size=8, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, 'AgriLink MVP | Farmer-to-Buyer Agricultural Marketplace', size=8, color=MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(title, 'AgriLink', size=28, bold=True, color=INK)
    subtitle = doc.add_paragraph(style='Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, 'Complete Project Documentation | AI-Powered Farmer-to-Buyer Marketplace', size=13, color=MUTED)
    doc.add_paragraph()
    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2700, 6660])
    for row, left, right in [
        (0, 'Document purpose', 'Current technical and functional documentation for the AgriLink MVP'),
        (1, 'Architecture', 'MERN: React/Vite, Node.js/Express, MongoDB/Mongoose'),
        (2, 'Prepared', date.today().strftime('%d %B %Y')),
    ]:
        set_cell_shading(meta.cell(row, 0), LIGHT_BLUE)
        add_text(meta.cell(row, 0).paragraphs[0], left, bold=True, color=INK)
        add_text(meta.cell(row, 1).paragraphs[0], right)
    doc.add_paragraph()
    add_note(doc, 'Document basis.', 'This report follows the structure of the supplied reference document, but all descriptions, flows, entities, APIs and limitations have been adapted to the actual AgriLink codebase. It is therefore an implementation document, not a generic project proposal.')

    add_heading(doc, '1. Executive Summary')
    add_paragraph(doc, 'AgriLink is a full-stack agricultural marketplace that connects farmers directly with consumers, retailers and wholesalers. It reduces the gap between farm supply and buyer demand by allowing verified farmers to publish produce listings, buyers to discover verified products and place orders, and administrators to govern user and product verification.')
    add_paragraph(doc, 'The application implements the complete marketplace path: farmer listing creation, administrator review, public product discovery, buyer cart and checkout, server-side stock reservation, farmer fulfilment and role-specific dashboard analytics. The current MVP also includes wishlists, verified-purchase reviews, public farmer storefronts, multi-quantity pricing support and a local AI assistant powered by Ollama.')
    add_table(doc, ['Stakeholder', 'Primary value delivered'], [
        ('Farmers', 'Create listings, attach images, maintain inventory, fulfil orders and monitor sales.'),
        ('Buyers', 'Find verified produce, save products, manage cart quantities, order securely and review completed purchases.'),
        ('Administrators', 'Verify farmers/products, view marketplace activity and manage operational visibility.'),
    ], [2500, 6860])

    add_heading(doc, '2. Problem Statement and Objectives')
    add_paragraph(doc, 'Small and medium farmers often depend on fragmented market channels, have limited product visibility and receive delayed feedback about demand. Buyers, in turn, need an easier way to find trustworthy local produce with clear availability, pricing and seller information. AgriLink addresses these gaps with a controlled direct-marketplace workflow.')
    add_heading(doc, '2.1 Objectives', 2)
    add_bullets(doc, [
        'Give farmers a simple digital channel for publishing agricultural products and maintaining stock.',
        'Give consumers, retailers and wholesalers a searchable marketplace of active, verified products.',
        'Protect marketplace quality through separate user and product verification states.',
        'Make ordering reliable by calculating price and inventory changes on the server, not in the browser.',
        'Provide role-based dashboards for operational decisions and marketplace oversight.',
        'Offer contextual AI assistance without exposing data outside the user’s permitted role.',
    ])
    add_heading(doc, '2.2 MVP Scope', 2)
    add_table(doc, ['Included in the current MVP', 'Not yet implemented'], [
        ('JWT sign-up/login, roles and verification', 'Online payment gateway and settlement'),
        ('Product, cart, order, wishlist and review workflows', 'Delivery-partner or route optimisation integration'),
        ('Image uploads using local development storage or Cloudinary', 'Image management UI for deleting/reordering individual images'),
        ('Dashboard polling every 15 seconds', 'WebSocket push updates'),
        ('Ollama-based assistant and product recommendations', 'Hosted AI model / multilingual voice assistant'),
    ], [4680, 4680])

    add_heading(doc, '3. System Requirements')
    add_heading(doc, '3.1 Software Requirements', 2)
    add_table(doc, ['Component', 'Requirement / purpose'], [
        ('Operating system', 'Windows 10/11, macOS or Linux'),
        ('Node.js', 'Node.js 18+ recommended; the current development environment has Node.js 24'),
        ('Package manager', 'npm'),
        ('Database', 'MongoDB running locally or available through MongoDB Atlas'),
        ('Browser', 'Current Chrome, Edge or Firefox for UI testing'),
        ('Optional AI runtime', 'Ollama with phi:latest for the AI assistant'),
        ('Optional media service', 'Cloudinary credentials for production image storage'),
        ('API testing', 'Postman collection included in the repository'),
    ], [2500, 6860])
    add_heading(doc, '3.2 Hardware Guidance', 2)
    add_bullets(doc, [
        'Processor: Intel Core i5 / AMD Ryzen 5 or equivalent for comfortable local development.',
        'Memory: 8 GB minimum; 16 GB recommended when running MongoDB, two development servers, an IDE and browser tools together.',
        'Storage: at least 2 GB free for dependencies, local MongoDB data and uploaded development images.',
        'Display: 1366 x 768 minimum; higher resolution improves dashboard and responsive-layout testing.',
    ])

    add_heading(doc, '4. Technology Stack')
    add_table(doc, ['Layer', 'Technology', 'Responsibility'], [
        ('Frontend', 'React 18 + Vite', 'Single-page user interface, routes, stateful pages and API calls.'),
        ('Styling', 'Tailwind CSS + Lucide React', 'Responsive UI, utility styling and interface icons.'),
        ('Backend', 'Node.js + Express', 'REST API, middleware, authorisation and business rules.'),
        ('Database', 'MongoDB + Mongoose', 'Document data model, validation, indexes and references.'),
        ('Security', 'JWT + bcryptjs + Helmet + CORS + rate limiting', 'Session authentication, password hashing and request protection.'),
        ('Uploads', 'Cloudinary or local API storage', 'Cloud media in production, local files for development without credentials.'),
        ('AI', 'Ollama + phi:latest', 'Local, role-aware assistant and product assistance.'),
    ], [1500, 2700, 5160])

    add_heading(doc, '5. Technical Architecture')
    add_paragraph(doc, 'AgriLink follows a client-server architecture. The React application provides the presentation layer. It calls a REST API hosted by Express. Controllers apply authorisation and business rules, then interact with Mongoose models stored in MongoDB. Uploaded images are stored either in Cloudinary or, for localhost development, in a backend uploads directory exposed by Express.')
    add_code(doc, [
        'User browser',
        '    ↓ HTTPS / HTTP JSON requests',
        'React + Vite frontend (localhost:5173)',
        '    ↓ Axios with optional JWT bearer token',
        'Express API (localhost:5000/api)',
        '    ↓ Controllers, middleware, validation, business rules',
        'MongoDB via Mongoose        Cloudinary or backend/uploads',
        '    ↓',
        'Role-specific dashboard, marketplace, orders and analytics',
    ])
    add_heading(doc, '5.1 MVC and Layer Responsibilities', 2)
    add_table(doc, ['Layer', 'AgriLink implementation'], [
        ('Model', 'Mongoose models define User, Product, Cart, Order, Wishlist and Review data with validation and indexes.'),
        ('Controller', 'Controllers perform validation, enforce ownership/roles, compute order prices, reserve stock and shape responses.'),
        ('Routes / API view', 'Express routes expose REST resources and compose middleware such as protect, optionalAuth and allowRoles.'),
        ('Frontend view', 'React pages and shared components present flows for public visitors, buyers, farmers and administrators.'),
    ], [2500, 6860])

    add_heading(doc, '6. Roles, Responsibilities and User Flows')
    add_heading(doc, '6.1 Buyer', 2)
    add_bullets(doc, [
        'Registers as a buyer and logs in through a JWT-protected session.',
        'Browses verified, active marketplace listings with search, category, price, stock and product-type filters.',
        'Adds products to a cart or wishlist, then submits an order with a complete delivery address.',
        'Views order history, cancels permitted orders and leaves reviews only after completed purchases.',
    ])
    add_heading(doc, '6.2 Farmer', 2)
    add_bullets(doc, [
        'Registers as a farmer, maintains profile information and awaits user verification when required.',
        'Creates, edits and deactivates product listings; edited products are submitted for re-verification.',
        'Uploads images while editing. Local images work without Cloudinary on localhost; Cloudinary is used when configured.',
        'Receives buyer orders, progresses order status and monitors inventory, sales and low-stock signals.',
    ])
    add_heading(doc, '6.3 Administrator', 2)
    add_bullets(doc, [
        'Reviews and verifies or rejects user accounts and farmer product submissions.',
        'Views marketplace-wide metrics, user records and orders.',
        'Can remove reviews when moderation is required.',
    ])
    add_heading(doc, '6.4 Marketplace Flow', 2)
    add_numbered(doc, [
        'Farmer creates a product; it starts in the pending verification state.',
        'Administrator verifies the product; only active and verified products are exposed publicly.',
        'Buyer searches, opens the product, adds a quantity to cart and completes checkout.',
        'The API validates stock and minimum order quantity, calculates the final price and reserves stock atomically.',
        'Farmer moves the order through confirmed, processing and completed states; buyer may cancel only when the state transition allows it.',
        'Dashboards refresh automatically and report updated product, order and inventory metrics.',
    ])

    add_heading(doc, '7. Functional Modules')
    add_table(doc, ['Module', 'Implemented behaviour'], [
        ('Authentication and profile', 'Registration, login, JWT issuance, password hashing, profile editing and role control.'),
        ('Marketplace', 'Public browsing of verified active products, search, filter, sort, pagination and public farmer information.'),
        ('Product management', 'Farmer CRUD, inventory, agricultural metadata, wholesale pricing fields, images and soft deactivation.'),
        ('Cart and checkout', 'Buyer-only cart, quantity validation, one-farmer-per-order rule, delivery address and server-calculated totals.'),
        ('Order fulfilment', 'Buyer/farmer/admin access rules, transition validation, history and inventory restoration on cancellation/rejection.'),
        ('Dashboard analytics', 'Role-specific metrics; product/order/verification workspace polling every 15 seconds.'),
        ('Wishlist and reviews', 'Buyer wishlist; verified-purchase ratings and review moderation.'),
        ('Farmer storefront', 'Public seller details and their active, verified product listings.'),
        ('AI assistant', 'Authenticated role-aware product lookup, recommendations and local Ollama chat context.'),
    ], [2700, 6660])

    add_heading(doc, '8. Data Model and Entity Relationships')
    add_paragraph(doc, 'MongoDB stores related marketplace data using ObjectId references. The following conceptual relationships correspond to the current Mongoose models.')
    add_code(doc, [
        'User (farmer) 1 ─── * Product',
        'User (buyer)  1 ─── 1 Cart',
        'Cart          1 ─── * Cart items ─── 1 Product',
        'User (buyer)  1 ─── * Order',
        'User (farmer) 1 ─── * Order',
        'Order         1 ─── * Order items ─── 1 Product snapshot',
        'User (buyer)  1 ─── 1 Wishlist ─── * Product',
        'Product       1 ─── * Review; User (buyer) 1 ─── * Review',
    ])
    add_table(doc, ['Entity', 'Key data held'], [
        ('User', 'Name, email, password hash, role, verification status, contact/address details, active status.'),
        ('Product', 'Farmer reference, name, description, category, price/unit, stock, images, location, verification and active status.'),
        ('Order', 'Buyer/farmer references, item price snapshots, quantities, total, commission, delivery address and status history.'),
        ('Cart', 'Buyer reference and an array of product/quantity line items.'),
        ('Wishlist', 'Buyer reference and a de-duplicated product collection.'),
        ('Review', 'Product/buyer references, rating, written feedback and verified-purchase marker.'),
    ], [2200, 7160])
    add_heading(doc, '8.1 Product Lifecycle', 2)
    add_paragraph(doc, 'A product has two separate state concerns. Verification status is pending, verified or rejected; marketplace status is active or inactive. Public browsing requires both active and verified. Deactivation is a soft delete: it preserves order history while removing the listing from public and active farmer-management results.')

    add_heading(doc, '9. API Overview')
    add_paragraph(doc, 'All API routes are served below /api. Protected endpoints require the Authorization: Bearer <token> header. The routes below are the main public contract; the included Postman collection can be used for manual API checks.')
    add_table(doc, ['Resource', 'Representative endpoints', 'Access'], [
        ('Health', 'GET /health', 'Public'),
        ('Auth', 'POST /auth/register, POST /auth/login, GET /auth/me', 'Public / authenticated'),
        ('Users', 'PATCH /users/me, GET /users, PATCH /users/:id/verification', 'Authenticated / admin'),
        ('Products', 'GET /products, POST/PATCH/DELETE /products/:id, PATCH verification', 'Public read; farmer/admin writes'),
        ('Cart', 'GET /cart, POST/PATCH/DELETE cart items', 'Buyer'),
        ('Wishlist', 'GET/POST/DELETE /wishlist', 'Buyer'),
        ('Orders', 'GET/POST /orders, GET /orders/:id, PATCH status', 'Authenticated; role rules'),
        ('Reviews', 'GET/POST product reviews, PATCH/DELETE review', 'Public read; buyer/admin write rules'),
        ('Farmers', 'GET /farmers/:id', 'Public'),
        ('Analytics', 'GET /analytics/dashboard', 'Authenticated'),
        ('Uploads', 'POST /uploads/image', 'Authenticated'),
        ('AI', 'POST /ai/chat, GET /ai/recommendations', 'Authenticated'),
    ], [1600, 5000, 2760])

    add_heading(doc, '10. Security and Business Rules')
    add_bullets(doc, [
        'Passwords are stored with bcrypt hashes; plaintext passwords are never stored in the user model.',
        'JWT tokens are used to identify the signed-in user. Protect middleware reads the token before private routes run.',
        'Role middleware limits product management to farmers, cart/wishlist use to buyers and verification to administrators.',
        'Helmet, CORS configuration, rate limiting, MongoDB sanitisation and a central error handler provide baseline API protection.',
        'Only verified and active products are public. Product ownership is checked before a farmer can edit or deactivate a listing.',
        'Orders must contain products from one farmer, a complete delivery address and valid quantities. Product prices are recalculated by the server.',
        'Stock is reserved through conditional database updates. A failed reservation rolls back any earlier reservation within the same request.',
        'Buyers can review only completed purchases; review changes update the product rating aggregate.',
    ])

    add_heading(doc, '11. Image Upload Design')
    add_paragraph(doc, 'The upload endpoint accepts image files only and applies a 5 MB limit. It behaves differently by environment so that frontend work can be tested before external cloud configuration is available.')
    add_table(doc, ['Environment condition', 'Storage and result'], [
        ('Cloudinary credentials configured', 'The API streams the image to the agrilink Cloudinary folder and returns its secure URL/public ID.'),
        ('No Cloudinary credentials (local development)', 'The API stores the image in backend/uploads and returns a URL served from the Express API.'),
        ('Existing product editing', 'The dashboard appends a newly uploaded image object to the existing images array; the user then saves the product changes.'),
    ], [3500, 5860])
    add_note(doc, 'Operational note.', 'Local uploaded files are ignored by Git. For deployment, use Cloudinary or another durable object store because host-local files may not persist across deployments.')

    add_heading(doc, '12. AI Assistant')
    add_paragraph(doc, 'AgriLink AI is designed to run locally with Ollama rather than a paid external AI API. The assistant retrieves only data appropriate to the current role: buyers receive recent-order context and recommendations, farmers receive inventory and sales context, and administrators receive marketplace aggregates. Product-search intent returns deterministic product data from MongoDB before any model-generated response is used.')
    add_bullets(doc, [
        'Set OLLAMA_URL (default http://127.0.0.1:11434) and OLLAMA_MODEL (default phi:latest) in backend/.env.',
        'Start Ollama and install the selected model before using chat.',
        'If Ollama is unavailable, the API returns a clear 503 error rather than inventing an answer.',
        'The prompt restricts the model to supplied role-authorized data and asks for concise responses.',
    ])

    add_heading(doc, '13. Local Setup and Execution')
    add_heading(doc, '13.1 Environment Variables', 2)
    add_table(doc, ['Variable', 'Use'], [
        ('MONGODB_URI', 'Required MongoDB connection string.'),
        ('JWT_SECRET', 'Required long secret used to sign access tokens.'),
        ('PORT', 'Optional API port; defaults to 5000.'),
        ('FRONTEND_URL', 'Allowed frontend origin(s) for CORS.'),
        ('VITE_API_URL', 'Optional frontend API base URL; defaults to http://localhost:5000/api.'),
        ('CLOUDINARY_CLOUD_NAME / API_KEY / API_SECRET', 'Optional production image storage configuration.'),
        ('OLLAMA_URL / OLLAMA_MODEL', 'Optional local AI configuration.'),
        ('PLATFORM_COMMISSION_PERCENT', 'Optional order commission percentage; defaults to 2.'),
    ], [3700, 5660])
    add_heading(doc, '13.2 Start Commands', 2)
    add_code(doc, [
        '# From the repository root',
        'npm.cmd install',
        'npm.cmd run seed',
        'npm.cmd run dev',
        '',
        '# Or use two terminals',
        'npm.cmd run dev --workspace backend',
        'npm.cmd run dev --workspace frontend',
    ])
    add_paragraph(doc, 'Open http://localhost:5173 for the frontend. The backend health endpoint is http://localhost:5000/api/health. MongoDB must be running before the API starts.')
    add_heading(doc, '13.3 Demo Accounts', 2)
    add_table(doc, ['Role', 'Email', 'Password'], [
        ('Administrator', 'admin@agrilink.in', 'Admin@123'),
        ('Farmer', 'ramesh@agrilink.in', 'Farmer@123'),
        ('Buyer', 'priya@agrilink.in', 'Buyer@123'),
    ], [2300, 4000, 3060])

    add_heading(doc, '14. Testing Strategy and Manual Acceptance Checks')
    add_paragraph(doc, 'The repository includes a Postman collection and seeded data for functional checks. The backend test command currently completes successfully but has no automated test suites, so manual acceptance testing remains important.')
    add_table(doc, ['Scenario', 'Expected result'], [
        ('Farmer creates product', 'Product is created as pending and is visible to the administrator for review.'),
        ('Administrator verifies product', 'Product becomes visible in public marketplace results when active.'),
        ('Farmer edits/adds image', 'New image is appended; save sends the listing for re-verification.'),
        ('Farmer deactivates product', 'Product disappears from active dashboard and public marketplace; order history remains intact.'),
        ('Buyer checkout', 'API validates stock, MOQ and one-farmer rule; order total is server-calculated and stock reduces.'),
        ('Order cancellation/rejection', 'Allowed transition restores inventory and records the new status.'),
        ('Dashboard update', 'Products, orders and analytics refresh automatically within approximately 15 seconds.'),
        ('AI unavailable', 'UI/API reports a clear service-unavailable state if Ollama is stopped.'),
    ], [3100, 6260])

    add_heading(doc, '15. Current Limitations and Recommended Roadmap')
    add_table(doc, ['Priority', 'Recommended enhancement', 'Reason'], [
        ('High', 'Add automated API and frontend tests', 'Prevents regressions in stock, order transitions and dashboard behaviour.'),
        ('High', 'Add payment and settlement integration', 'Completes the commercial transaction lifecycle.'),
        ('High', 'Use managed cloud media and database backups in production', 'Ensures uploaded media and data survive deployments.'),
        ('Medium', 'Replace polling with WebSockets/SSE where needed', 'Provides immediate operational updates at scale.'),
        ('Medium', 'Add image gallery management', 'Lets farmers preview, remove, reorder and select a cover image.'),
        ('Medium', 'Add delivery tracking and notifications', 'Improves buyer confidence after checkout.'),
        ('Low', 'Expand AI with regional languages and voice input', 'Improves accessibility for wider farming communities.'),
    ], [1100, 4100, 4160])

    add_heading(doc, '16. Repository Structure')
    add_code(doc, [
        'agrilink/',
        '├── frontend/                 React + Vite + Tailwind client',
        '│   └── src/pages/            Marketplace, dashboard, wishlist and detail pages',
        '├── backend/',
        '│   └── src/',
        '│       ├── models/           Mongoose schemas',
        '│       ├── controllers/      Marketplace business logic',
        '│       ├── routes/           REST endpoint definitions',
        '│       ├── middleware/       Auth and error handling',
        '│       ├── ai/               Ollama integration',
        '│       └── uploads/          Local development images (ignored by Git)',
        '├── postman_collection.json   Importable API tests',
        '└── package.json              Workspaces and root scripts',
    ])

    add_heading(doc, '17. Conclusion')
    add_paragraph(doc, 'AgriLink is a functioning MERN marketplace MVP with role-based access, controlled product verification, inventory-aware ordering, seller tooling and buyer discovery. Its most important design choice is that business-critical rules—availability, price, ownership, product visibility and order transitions—are enforced by the API rather than trusted to the frontend. This gives the project a sound base for continued work on payments, delivery workflows, automated testing and production deployment.')

    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
